"""
Conversion de la sortie JSON de l'agent "Rédacteur final" en document Word (.docx),
via python-docx.
"""

import json
import re
from io import BytesIO

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _strip_fences(raw: str) -> str:
    # Les modèles peuvent entourer le JSON de balises Markdown ; on les retire
    # avant tout passage au parseur strict.
    text = raw.strip()
    text = re.sub(r"^```(json)?", "", text.strip(), flags=re.IGNORECASE).strip()
    text = re.sub(r"```$", "", text.strip()).strip()
    return text


def parse_redaction_json(raw: str) -> dict:
    """
    Extrait un objet JSON depuis la sortie brute du LLM, en étant tolérant :
    - retire les éventuelles balises ```json ... ```
    - si du texte parasite entoure le JSON, isole le premier bloc {...}
    """
    text = _strip_fences(raw)

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # Fallback : isoler le premier bloc { ... } équilibré
    start = text.find("{")
    if start == -1:
        raise ValueError("Aucun objet JSON trouvé dans la sortie de l'agent rédacteur.")

    depth = 0
    for i in range(start, len(text)):
        if text[i] == "{":
            depth += 1
        elif text[i] == "}":
            depth -= 1
            if depth == 0:
                candidate = text[start : i + 1]
                return json.loads(candidate)

    raise ValueError("JSON mal formé dans la sortie de l'agent rédacteur.")


def diagnose_json_error(raw: str) -> str | None:
    """
    Retourne None si le JSON (après nettoyage des balises ```) est valide.
    Sinon, retourne un rapport détaillé — message d'erreur + quelques lignes
    de contexte autour de la position fautive avec un repère '^' — pensé pour
    être injecté tel quel dans le prompt d'un agent correcteur.
    """
    text = _strip_fences(raw)
    try:
        json.loads(text)
        return None
    except json.JSONDecodeError as e:
        lines = text.split("\n")
        lineno = e.lineno  # 1-indexé
        colno = e.colno  # 1-indexé

        start = max(1, lineno - 2)
        end = min(len(lines), lineno + 2)
        context_rows = []
        for i in range(start, end + 1):
            marker = ">>" if i == lineno else "  "
            context_rows.append(f"{marker} {i:>3}: {lines[i - 1]}")
            if i == lineno:
                # Repère '^' aligné sous le caractère fautif
                pointer = " " * (len(f"{marker} {i:>3}: ") + colno - 1) + "^"
                context_rows.append(pointer)

        return (
            f"Erreur JSON : {e.msg}\n"
            f"Position : ligne {lineno}, colonne {colno} (caractère {e.pos})\n\n"
            "Contexte (la ligne '>>' est la ligne fautive, '^' pointe le caractère problématique) :\n"
            + "\n".join(context_rows)
        )


def build_docx(data: dict) -> BytesIO:
    """
    Construit un document Word à partir du dict structuré (issu de parse_redaction_json)
    et le retourne comme buffer en mémoire, prêt pour st.download_button.
    """
    doc = Document()

    # Style de base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    titre = data.get("titre") or "Compte-rendu de réunion"
    date = data.get("date") or ""

    heading = doc.add_heading(titre, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if date:
        p = doc.add_paragraph(date)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True

    # Participants / absents
    doc.add_heading("Participants", level=1)
    participants = data.get("participants") or []
    if participants:
        for name in participants:
            doc.add_paragraph(str(name), style="List Bullet")
    else:
        doc.add_paragraph("Non précisé.")

    absents = data.get("absents") or []
    if absents:
        doc.add_heading("Absents / excusés", level=2)
        for name in absents:
            doc.add_paragraph(str(name), style="List Bullet")

    # Objectif
    doc.add_heading("Objectif de la réunion", level=1)
    doc.add_paragraph(data.get("objectif") or "Non précisé.")

    # Points clés
    doc.add_heading("Points clés abordés", level=1)
    points_cles = data.get("points_cles") or []
    if points_cles:
        for pc in points_cles:
            doc.add_paragraph(str(pc), style="List Bullet")
    else:
        doc.add_paragraph("Non précisé.")

    # Outils & chiffres associés
    doc.add_heading("Outils & chiffres associés", level=1)
    outils_et_chiffres = data.get("outils_et_chiffres") or []
    if outils_et_chiffres:
        for item in outils_et_chiffres:
            outil_name = str(item.get("outil", "")) if isinstance(item, dict) else str(item)
            p = doc.add_paragraph()
            p.add_run(outil_name).bold = True

            chiffres_associes = item.get("chiffres_associes") if isinstance(item, dict) else None
            chiffres_associes = chiffres_associes or []
            if chiffres_associes:
                for c in chiffres_associes:
                    doc.add_paragraph(str(c), style="List Bullet")
            else:
                doc.add_paragraph("Aucun chiffre associé.", style="List Bullet")
    else:
        doc.add_paragraph("Aucun outil mentionné.")

    autres_chiffres = data.get("autres_chiffres") or []
    if autres_chiffres:
        doc.add_heading("Autres chiffres clés", level=2)
        for c in autres_chiffres:
            doc.add_paragraph(str(c), style="List Bullet")

    # Décisions
    doc.add_heading("Décisions prises", level=1)
    decisions = data.get("decisions") or []
    if decisions:
        for d in decisions:
            doc.add_paragraph(str(d), style="List Bullet")
    else:
        doc.add_paragraph("Aucune décision actée.")

    # Actions
    doc.add_heading("Actions à faire", level=1)
    actions = data.get("actions") or []
    if actions:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Action"
        hdr[1].text = "Responsable"
        hdr[2].text = "Échéance"
        for a in actions:
            row = table.add_row().cells
            row[0].text = str(a.get("action", ""))
            row[1].text = str(a.get("responsable", "") or "Non précisé")
            row[2].text = str(a.get("echeance", "") or "Non précisée")
    else:
        doc.add_paragraph("Aucune action identifiée.")

    # Points de blocage
    doc.add_heading("Points de blocage / questions ouvertes", level=1)
    points = data.get("points_de_blocage") or []
    if points:
        for pt in points:
            doc.add_paragraph(str(pt), style="List Bullet")
    else:
        doc.add_paragraph("Aucun point de blocage identifié.")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer