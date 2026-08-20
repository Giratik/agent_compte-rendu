# frontend/transcriber_bundle/transcriber_page.py

import streamlit as st
import requests
import time
import os
import uuid

from utility.session_state_central_cr import SK, get, set as ss_set, init_session_state

API_URL = os.environ.get("BACKEND_URL", "http://backend:8000")
WHISPER_MODEL = os.environ.get("WHISPER_MODEL", "large-v3")

def render_transcriber():
    """Plugin UI autonome pour la transcription"""
    st.subheader("🎙️ Module de Transcription")
    
    input_mode = st.radio(
        "Comment voulez-vous fournir le texte ?",
        ["Fichier Audio/Vidéo", "Fichier Texte existant (.txt/.docx)", "Coller du texte"],
        horizontal=True,
        key="transcriber_input_mode"
    )

    if input_mode == "Fichier Audio/Vidéo":
        uploaded_file = st.file_uploader(
            "Déposez votre fichier audio ou vidéo",
            type=["wav", "mp3", "m4a", "flac", "mp4", "avi", "mov", "mkv", "webm"],
            key="transcriber_uploader"
        )

        en_file_attente = get(SK.TRANSCRIPTION_QUEUE_TOKEN) is not None
        
        # --- 2. TRAITEMENT DE LA TRANSCRIPTION ---
        # Déplacé AVANT le rendu du bouton : plus besoin de st.rerun() destructeur à la fin.
        if get(SK.IS_TRANSCRIBING) and uploaded_file is not None:
            
            msg_spinner = (
                "Envoi au serveur et transcription en cours (cela peut prendre plusieurs minutes)..."
                if en_file_attente
                else "Demande d'une place dans la file d'attente..."
            )

            with st.spinner(msg_spinner):
                start_time = time.time()
                files = {"file": (uploaded_file.name, uploaded_file.getvalue(), uploaded_file.type)}
                data = {"model_choice": WHISPER_MODEL, "session_id": get(SK.SESSION_ID)}

                if get(SK.TRANSCRIPTION_QUEUE_TOKEN):
                    data["queue_token"] = get(SK.TRANSCRIPTION_QUEUE_TOKEN)

                try:
                    response = requests.post(f"{API_URL}/transcribe/processing_audio", files=files, data=data)
                    response.raise_for_status()
                    response_data = response.json()

                    if response_data.get("status") in ["queued", "waiting"]:
                        ss_set(SK.TRANSCRIPTION_QUEUE_TOKEN, response_data.get(
                            "queue_token", get(SK.TRANSCRIPTION_QUEUE_TOKEN)
                        ))
                        ss_set(SK.TRANSCRIPTION_QUEUE_POSITION, response_data["position"])
                        ss_set(SK.IS_TRANSCRIBING, False)
                        en_file_attente = True
                        st.rerun()

                    elif "transcript" in response_data:
                        ss_set(SK.TRANSCRIPTION_QUEUE_TOKEN, None)
                        ss_set(SK.TRANSCRIPTION_QUEUE_POSITION, None)
                        ss_set(SK.TRANSCRIPT_TEXT, response_data["transcript"])
                        ss_set(SK.TOKEN_COUNT, len(response_data["transcript"].split()))
                        
                        execution_time = time.time() - start_time
                        st.success(f"Transcription terminée en {int(execution_time // 60)} min !")
                        st.toast('🎙️ La transcription audio est terminée !', icon='✅')
                        if get(SK.TOKEN_COUNT):
                            st.info(f"📊 Nombre de tokens dans le transcript : {get(SK.TOKEN_COUNT)}")
                        
                        ss_set(SK.IS_TRANSCRIBING, False)

                    else:
                        st.error(f"Réponse inattendue: {response_data}")
                        ss_set(SK.TRANSCRIPTION_QUEUE_TOKEN, None)
                        ss_set(SK.TRANSCRIPTION_QUEUE_POSITION, None)
                        ss_set(SK.IS_TRANSCRIBING, False)

                except Exception as e:
                    st.error(f"Erreur de communication : {e}")
                    ss_set(SK.TRANSCRIPTION_QUEUE_TOKEN, None)
                    ss_set(SK.TRANSCRIPTION_QUEUE_POSITION, None)
                    ss_set(SK.IS_TRANSCRIBING, False)

        # --- 3. GESTION DU BOUTON PRINCIPAL ---
        bouton_disabled = (uploaded_file is None) or get(SK.IS_TRANSCRIBING) or en_file_attente

        if en_file_attente and get(SK.TRANSCRIPTION_QUEUE_POSITION) is not None and get(SK.TRANSCRIPTION_QUEUE_POSITION) > 1:
            label_bouton = f"🕒 En attente (Position {get(SK.TRANSCRIPTION_QUEUE_POSITION)})"
        elif get(SK.IS_TRANSCRIBING):
            label_bouton = "⌛ Transcription en cours..."
        else:
            label_bouton = "Lancer la transcription"

        if st.button(label_bouton, disabled=bouton_disabled, key="btn_run_transcription"):
            ss_set(SK.IS_TRANSCRIBING, True)
            st.rerun()

        # --- 4. POLLING DE LA FILE D'ATTENTE ---
        if en_file_attente and not get(SK.IS_TRANSCRIBING):
            if get(SK.TRANSCRIPTION_QUEUE_POSITION) is not None and get(SK.TRANSCRIPTION_QUEUE_POSITION) > 1:
                st.info(f"⏳ Vous êtes en position **{get(SK.TRANSCRIPTION_QUEUE_POSITION)}** dans la file d'attente. Veuillez patienter...")

                try:
                    response = requests.get(f"{API_URL}/transcribe/queue-position", params={
                        "queue_token": get(SK.TRANSCRIPTION_QUEUE_TOKEN)
                    })

                    if response.status_code == 200:
                        position_data = response.json()
                        if position_data.get("status") == "waiting":
                            ss_set(SK.TRANSCRIPTION_QUEUE_POSITION, position_data.get("position"))
                        elif position_data.get("status") == "processing":
                            ss_set(SK.TRANSCRIPTION_QUEUE_POSITION, 1)
                    else:
                        time.sleep(3)
                        st.rerun()

                    if get(SK.TRANSCRIPTION_QUEUE_POSITION) == 1:
                        st.success("C'est votre tour ! Démarrage imminent...")
                        time.sleep(1)
                        ss_set(SK.IS_TRANSCRIBING, True)
                        st.rerun()
                    else:
                        time.sleep(2 if get(SK.TRANSCRIPTION_QUEUE_POSITION) <= 2 else 5)
                        st.rerun()

                except Exception as e:
                    time.sleep(5)
                    st.rerun()

            elif get(SK.TRANSCRIPTION_QUEUE_POSITION) == 1:
                ss_set(SK.IS_TRANSCRIBING, True)
                st.rerun()

    elif input_mode == "Fichier Texte existant (.txt/.docx)":
        uploaded_text = st.file_uploader(
            "Déposez votre fichier texte",
            type=["txt", "docx"],
            key="txt_uploader"
        )
        if uploaded_text is not None:
            if st.button("Charger ce fichier", key="btn_load_txt"):
                file_content = None
                try:
                    if uploaded_text.name.lower().endswith(".docx"):
                        from docx import Document
                        import io

                        doc = Document(io.BytesIO(uploaded_text.getvalue()))
                        parts = []
                        for para in doc.paragraphs:
                            if para.text.strip():
                                parts.append(para.text)
                        for table in doc.tables:
                            for row in table.rows:
                                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                                if row_text.strip(" |"):
                                    parts.append(row_text)
                        file_content = "\n".join(parts)

                        if not file_content.strip():
                            st.warning("⚠️ Le fichier Word semble vide ou ne contient pas de texte extractible.")
                            file_content = None
                    else:
                        try:
                            file_content = uploaded_text.getvalue().decode("utf-8")
                        except UnicodeDecodeError:
                            file_content = uploaded_text.getvalue().decode("latin-1")

                except Exception as e:
                    st.error(f"❌ Erreur lors du chargement du fichier : {str(e)}")
                    file_content = None

                if file_content is not None:
                    ss_set(SK.TRANSCRIPT_TEXT, file_content)
                    token_count = len(file_content.split())
                    ss_set(SK.TOKEN_COUNT, token_count)
                    st.success("✅ Fichier chargé avec succès !")
                    st.info(f"📊 Nombre de tokens dans le transcript : {token_count}")

    elif input_mode == "Coller du texte":
        pasted_text = st.text_area(
            "Collez votre texte ici",
            height=300,
            key="pasted_text_area",
            placeholder="Collez votre transcript"
        )
        if st.button("Charger ce texte", key="btn_load_pasted"):
            if pasted_text.strip():
                ss_set(SK.TRANSCRIPT_TEXT, pasted_text)
                token_count = len(pasted_text.split())
                ss_set(SK.TOKEN_COUNT, token_count)
                st.success("✅ Texte chargé avec succès !")
                st.info(f"📊 Nombre de tokens dans le transcript : {token_count}")
            else:
                st.warning("⚠️ Le champ de texte est vide.")

    if get(SK.TRANSCRIPT_TEXT):
        st.markdown("---")
        st.write("### Génération du Compte-rendu")

        transcript_editor = st.text_area("Texte à analyser", value=get(SK.TRANSCRIPT_TEXT), height=250)
        ss_set(SK.TRANSCRIPT_TEXT, transcript_editor)
